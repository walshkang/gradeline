from __future__ import annotations

import html
import io
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

import fitz

from .types import JsonDict, SubmissionUnit

try:
    import defusedxml.ElementTree as defused_ET
except ImportError:
    import xml.etree.ElementTree as defused_ET  # type: ignore[no-redef] # nosec B314


def discover_submission_units(submissions_dir: Path) -> list[SubmissionUnit]:
    units: list[SubmissionUnit] = []
    for folder in sorted([p for p in submissions_dir.iterdir() if p.is_dir() and not p.name.startswith(".")], key=lambda p: p.name.lower()):
        convert_non_pdf_files_to_pdf(folder)
        pdfs = sorted(folder.rglob("*.pdf"), key=lambda p: str(p.relative_to(folder)).lower())
        if not pdfs:
            continue
        token, student_name = parse_folder_name(folder.name)
        units.append(
            SubmissionUnit(
                folder_path=folder,
                folder_relpath=folder.relative_to(submissions_dir),
                folder_token=token,
                student_name=student_name,
                pdf_paths=pdfs,
            )
        )
    return units


def parse_folder_name(name: str) -> tuple[str, str]:
    parts = [segment.strip() for segment in name.split(" - ")]
    token = parts[0] if parts else name
    student_name = parts[1] if len(parts) > 1 else name
    return token, student_name


def normalize_name(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", value.lower())


def parse_index_html(index_html_path: Path) -> list[JsonDict]:
    if not index_html_path.exists():
        return []
    text = index_html_path.read_text(encoding="utf-8", errors="ignore")
    name_matches = list(re.finditer(r"<b>([^<]+,\s*[^<]+)</b>", text, flags=re.IGNORECASE))
    entries: list[JsonDict] = []
    for idx, match in enumerate(name_matches):
        person = html.unescape(match.group(1)).strip()
        start = match.end()
        end = name_matches[idx + 1].start() if idx + 1 < len(text) else len(text)
        chunk = text[start:end]

        for file_match in re.finditer(
            r"valign=top>([^<]+?)<p[^>]*><b>Comments:</b><br>(.*?)</td><td valign=top><b>Submitted:</b><br>([^<]+)</td>",
            chunk,
            flags=re.IGNORECASE | re.DOTALL,
        ):
            filename = html.unescape(file_match.group(1)).strip()
            comments_html = file_match.group(2)
            submitted = html.unescape(file_match.group(3)).strip()
            comments = _html_to_text(comments_html)
            entries.append(
                {
                    "student_name": person,
                    "submitted_filename": filename,
                    "submitted_at": submitted,
                    "comments": comments,
                }
            )
    return entries


def _html_to_text(raw: str) -> str:
    scrubbed = re.sub(r"<br\s*/?>", "\n", raw, flags=re.IGNORECASE)
    scrubbed = re.sub(r"<[^>]+>", "", scrubbed)
    return html.unescape(scrubbed).strip()


def convert_non_pdf_files_to_pdf(folder: Path) -> None:
    # Filter files in folder, excluding hidden files & HTML indexes
    all_files = [
        p for p in folder.iterdir()
        if p.is_file()
        and not p.name.startswith(".")
        and p.suffix.lower() != ".pdf"
        and p.name.lower() not in ("index.html", "index.htm", "index.txt")
    ]
    if not all_files:
        return

    # Process Images (PNG, JPG, JPEG, HEIC, HEIF, TIFF, WEBP)
    images = [p for p in all_files if p.suffix.lower() in (".png", ".jpg", ".jpeg", ".heic", ".heif", ".tiff", ".tif", ".webp")]
    if images:
        _convert_images_to_pdf(folder, images)

    # Process Word Documents (.docx, .doc)
    word_files = [p for p in all_files if p.suffix.lower() in (".docx", ".doc")]
    if word_files:
        _convert_word_to_pdf(folder, word_files)

    # Process Spreadsheets (.xlsx, .xls)
    excel_files = [p for p in all_files if p.suffix.lower() in (".xlsx", ".xls")]
    if excel_files:
        _convert_excel_to_pdf(folder, excel_files)


def _convert_images_to_pdf(folder: Path, images: list[Path]) -> None:
    try:
        import pillow_heif
        pillow_heif.register_heif_opener()
    except Exception:
        pass
    from PIL import Image

    doc = fitz.open()
    for img_path in sorted(images, key=lambda p: p.name.lower()):
        pdf_bytes = None
        # PyMuPDF fitz (standard raster formats)
        try:
            img_doc = fitz.open(img_path)
            pdf_bytes = img_doc.convert_to_pdf()
            img_doc.close()
        except Exception:
            pass

        # PIL / pillow_heif (HEIC, TIFF, WEBP, etc.)
        if not pdf_bytes:
            try:
                with Image.open(img_path) as im:
                    rgb_im = im.convert("RGB")
                    buf = io.BytesIO()
                    rgb_im.save(buf, format="PDF")
                    pdf_bytes = buf.getvalue()
            except Exception:
                pass

        # macOS sips CLI fallback
        if not pdf_bytes:
            try:
                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                    tmp_pdf = Path(tmp.name)
                res = subprocess.run(
                    ["sips", "-s", "format", "pdf", str(img_path), "--out", str(tmp_pdf)],
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    timeout=10,
                )
                if res.returncode == 0 and tmp_pdf.exists():
                    pdf_bytes = tmp_pdf.read_bytes()
                    tmp_pdf.unlink(missing_ok=True)
            except Exception:
                pass

        if pdf_bytes:
            try:
                doc.insert_pdf(fitz.open("pdf", pdf_bytes))
            except Exception:
                pass

    if len(doc) > 0:
        output_pdf = folder / f"{folder.name}_images.pdf"
        if output_pdf.exists() and len(images) == 1:
            output_pdf = folder / f"{images[0].stem}_image.pdf"
        try:
            doc.save(output_pdf)
        except Exception:
            pass
    doc.close()


def _convert_word_to_pdf(folder: Path, word_files: list[Path]) -> None:
    for word_path in word_files:
        output_pdf = _determine_output_pdf_path(folder, word_path)
        if output_pdf.exists():
            continue

        if _run_soffice_convert(word_path, folder, output_pdf):
            continue

        lines: list[str] = []
        is_docx = word_path.suffix.lower() == ".docx"

        if is_docx:
            try:
                import docx
                doc = docx.Document(word_path)
                for p in doc.paragraphs:
                    if p.text.strip():
                        lines.append(p.text.strip())
                for t in doc.tables:
                    for r in t.rows:
                        row_txt = [c.text.strip() for c in r.cells if c.text.strip()]
                        if row_txt:
                            lines.append("  |  ".join(row_txt))
            except Exception:
                pass

        if not lines and is_docx:
            lines = _parse_docx_raw_xml(word_path)

        if not lines and not is_docx:
            lines = _parse_doc_raw_text(word_path)

        if lines:
            _render_text_lines_to_pdf(lines, output_pdf)


def _convert_excel_to_pdf(folder: Path, excel_files: list[Path]) -> None:
    for excel_path in excel_files:
        output_pdf = _determine_output_pdf_path(folder, excel_path)
        if output_pdf.exists():
            continue

        if _run_soffice_convert(excel_path, folder, output_pdf):
            continue

        lines: list[str] = []
        is_xlsx = excel_path.suffix.lower() == ".xlsx"

        if is_xlsx:
            try:
                import openpyxl
                wb = openpyxl.load_workbook(excel_path, data_only=True)
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    lines.append(f"=== Sheet: {sheet_name} ===")
                    for row in ws.iter_rows(values_only=True):
                        cells = [str(c).strip() for c in row if c is not None and str(c).strip() != ""]
                        if cells:
                            lines.append("  |  ".join(cells))
            except Exception:
                pass
        else:
            try:
                import xlrd
                wb = xlrd.open_workbook(excel_path)
                for sheet_name in wb.sheet_names():
                    ws = wb.sheet_by_name(sheet_name)
                    lines.append(f"=== Sheet: {sheet_name} ===")
                    for r in range(ws.nrows):
                        row_vals = [str(ws.cell_value(r, c)).strip() for c in range(ws.ncols) if str(ws.cell_value(r, c)).strip() != ""]
                        if row_vals:
                            lines.append("  |  ".join(row_vals))
            except Exception:
                pass

        if not lines and is_xlsx:
            lines = _parse_xlsx_raw_xml(excel_path)

        if lines:
            _render_text_lines_to_pdf(lines, output_pdf)


def _determine_output_pdf_path(folder: Path, file_path: Path) -> Path:
    out_pdf = folder / f"{file_path.stem}.pdf"
    if out_pdf.exists() and out_pdf.name.lower() != file_path.name.lower() + ".pdf":
        out_pdf = folder / f"{file_path.stem}_{file_path.suffix.lstrip('.')}.pdf"
    return out_pdf


def _run_soffice_convert(input_path: Path, output_dir: Path, expected_pdf: Path) -> bool:
    soffice_bin = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice_bin:
        for possible_path in [
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            "/opt/homebrew/bin/soffice",
            "/usr/bin/soffice",
        ]:
            if Path(possible_path).exists() and Path(possible_path).is_file():
                soffice_bin = possible_path
                break
    if not soffice_bin:
        return False
    try:
        res = subprocess.run(
            [soffice_bin, "--headless", "--convert-to", "pdf", str(input_path), "--outdir", str(output_dir)],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=30,
            check=False,
        )
        soffice_pdf = output_dir / f"{input_path.stem}.pdf"
        if res.returncode == 0 and soffice_pdf.exists():
            if soffice_pdf != expected_pdf:
                soffice_pdf.rename(expected_pdf)
            return True
        return False
    except Exception:
        return False


def _parse_docx_raw_xml(docx_path: Path) -> list[str]:
    import zipfile
    lines: list[str] = []
    try:
        with zipfile.ZipFile(docx_path) as z:
            xml_content = z.read("word/document.xml")
            root = defused_ET.fromstring(xml_content)
            for paragraph in root.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"):
                para_text = []
                for run in paragraph.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
                    if run.text:
                        para_text.append(run.text)
                if para_text:
                    lines.append("".join(para_text))
    except Exception:
        pass
    return lines


def _parse_doc_raw_text(doc_path: Path) -> list[str]:
    lines: list[str] = []
    try:
        raw_bytes = doc_path.read_bytes()
        # Clean ASCII/UTF-8 readable chunks from legacy binary doc
        text = re.sub(r"[^\x20-\x7E\n\r\t]+", " ", raw_bytes.decode("latin1", errors="ignore"))
        chunks = [line.strip() for line in text.split("\n") if len(line.strip()) > 5]
        lines = chunks[:200]
    except Exception:
        pass
    return lines


def _parse_xlsx_raw_xml(xlsx_path: Path) -> list[str]:
    import zipfile
    lines: list[str] = []
    try:
        with zipfile.ZipFile(xlsx_path) as z:
            shared_strings = []
            try:
                strings_content = z.read("xl/sharedStrings.xml")
                strings_root = defused_ET.fromstring(strings_content)
                for t in strings_root.iter("{http://schemas.openxmlformats.org/spreadsheetml/2006/main}t"):
                    shared_strings.append(t.text or "")
            except KeyError:
                pass

            sheet_content = z.read("xl/worksheets/sheet1.xml")
            sheet_root = defused_ET.fromstring(sheet_content)
            rows = {}
            for row in sheet_root.iter("{http://schemas.openxmlformats.org/spreadsheetml/2006/main}row"):
                row_idx = int(row.get("r"))
                row_cells = []
                for cell in row.iter("{http://schemas.openxmlformats.org/spreadsheetml/2006/main}c"):
                    cell_ref = cell.get("r")
                    cell_type = cell.get("t")
                    val_el = cell.find("{http://schemas.openxmlformats.org/spreadsheetml/2006/main}v")
                    val = ""
                    if val_el is not None and val_el.text:
                        if cell_type == "s":
                            idx = int(val_el.text)
                            val = shared_strings[idx] if idx < len(shared_strings) else val_el.text
                        else:
                            val = val_el.text
                    row_cells.append((cell_ref, val))
                row_cells.sort(key=lambda x: x[0])
                rows[row_idx] = [val for ref, val in row_cells if val]
            for r_idx in sorted(rows.keys()):
                if rows[r_idx]:
                    lines.append("  |  ".join(rows[r_idx]))
    except Exception:
        pass
    return lines


def _render_text_lines_to_pdf(lines: list[str], output_pdf: Path) -> None:
    doc = fitz.open()
    page = doc.new_page()
    y = 50
    margin = 50
    line_height = 14
    for line in lines:
        if y > page.rect.height - 50:
            page = doc.new_page()
            y = 50
        max_chars = 95
        chunks = [line[i:i + max_chars] for i in range(0, len(line), max_chars)] or [""]
        for chunk in chunks:
            if y > page.rect.height - 50:
                page = doc.new_page()
                y = 50
            page.insert_text((margin, y), chunk, fontsize=9, fontname="helv")
            y += line_height
    doc.save(output_pdf)
    doc.close()
